"""
Success Criteria Validator for DocIngestionAgent.
Run from backend/ directory: python validate_sc.py
"""
import asyncio
import subprocess
import sys
import time
from io import BytesIO
from unittest.mock import AsyncMock, MagicMock, patch

results: list = []   # (name, passed, detail)


def record(name, passed, detail=""):
    results.append((name, passed, detail))
    icon = "PASS" if passed else "FAIL"
    print(f"  [{icon}]  {name}")
    if detail:
        print(f"           {detail}")


# ── Fixture helpers ────────────────────────────────────────────────────────────

def make_docx(n_para=10):
    from docx import Document
    doc = Document()
    doc.core_properties.title = "Benchmark"
    doc.core_properties.author = "Author"
    for i in range(n_para):
        if i % 8 == 0:
            doc.add_heading(f"Section {i // 8 + 1}", level=1)
        else:
            doc.add_paragraph(
                f"Paragraph {i}: Lorem ipsum dolor sit amet, consectetur "
                "adipiscing elit. Sed do eiusmod tempor incididunt. " * 3
            )
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_50page_docx():
    from docx import Document
    doc = Document()
    doc.core_properties.title = "50-Page Benchmark"
    doc.core_properties.author = "Benchmark Author"
    for section in range(50):
        doc.add_heading(f"Section {section + 1}: Research Overview", level=1)
        for p in range(8):
            doc.add_paragraph(
                f"S{section+1}/P{p+1}: "
                "This research explores novel methodologies for automated document "
                "processing using large language models and multi-agent frameworks. "
                "Results demonstrate superior performance across all benchmark tasks. "
                "The pipeline comprises five specialist agent nodes that collaborate "
                "asynchronously to produce publication-ready manuscripts. " * 2
            )
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_pdf():
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)
    for page in range(3):
        c.setFont("Helvetica-Bold", 14)
        c.drawString(72, 720, f"Section {page + 1}: Results and Discussion")
        c.setFont("Helvetica", 11)
        y = 690
        for line in range(20):
            c.drawString(
                72, y,
                f"Line {line}: The experimental results confirm our hypothesis "
                "with statistical significance p < 0.001."
            )
            y -= 18
        c.showPage()
    c.save()
    return buf.getvalue()


def make_text():
    body = "Natural language processing has advanced rapidly in recent years. " * 80
    return (
        "# Research Paper\n\n"
        "## Abstract\n\n"
        "This paper presents a comprehensive study of NLP pipelines (Jones 2022).\n\n"
        "## Introduction\n\n" + body + "\n\n"
        "## Methods\n\nWe used Python 3.12 for all experiments.\n\n"
        "## Results\n\nAccuracy improved by 12.3 percent over the baseline.\n\n"
        "## References\n\nJones, A. (2022). NLP in the wild. Journal of AI 10(2).\n"
    ).encode("utf-8")


# ── SC-1 ───────────────────────────────────────────────────────────────────────

async def sc1():
    print("\n[SC-1] All 3 formats produce IR objects that pass model_validate()")
    from app.agents.ingestion.docx_reader import DocxReader
    from app.agents.ingestion.pdf_reader import PdfReader
    from app.agents.ingestion.text_reader import TextReader
    from app.schemas.ir import IRSchema

    # .docx
    try:
        b = make_docx()
        with patch(
            "app.agents.ingestion.docx_reader.storage_service.upload_raw",
            new_callable=AsyncMock, return_value="s3://k",
        ):
            ir = await DocxReader().read(b, "sc1_docx")
        v = IRSchema.model_validate(ir.model_dump())
        ok = v.source_format == "docx" and len(v.elements) > 0
        record("SC-1a  .docx -> IRSchema.model_validate()", ok,
               f"source_format={v.source_format}, elements={len(v.elements)}, words={v.word_count}")
    except Exception as exc:
        record("SC-1a  .docx -> IRSchema.model_validate()", False, str(exc))

    # .pdf
    try:
        b = make_pdf()
        with (
            patch("app.agents.ingestion.pdf_reader.storage_service.upload_raw",
                  new_callable=AsyncMock, return_value="s3://k"),
            patch.object(PdfReader, "_is_scanned", return_value=False),
        ):
            ir = await PdfReader().read(b, "sc1_pdf")
        v = IRSchema.model_validate(ir.model_dump())
        ok = v.source_format == "pdf" and len(v.elements) > 0
        record("SC-1b  .pdf  -> IRSchema.model_validate()", ok,
               f"source_format={v.source_format}, elements={len(v.elements)}, words={v.word_count}")
    except Exception as exc:
        record("SC-1b  .pdf  -> IRSchema.model_validate()", False, str(exc))

    # .txt / .md
    try:
        b = make_text()
        ir = await TextReader().read(b, "sc1_txt")
        v = IRSchema.model_validate(ir.model_dump())
        ok = v.source_format == "txt" and len(v.elements) > 0
        record("SC-1c  .txt  -> IRSchema.model_validate()", ok,
               f"source_format={v.source_format}, elements={len(v.elements)}, words={v.word_count}")
    except Exception as exc:
        record("SC-1c  .txt  -> IRSchema.model_validate()", False, str(exc))


# ── SC-2 ───────────────────────────────────────────────────────────────────────

async def sc2():
    print("\n[SC-2] 50-page .docx completes in < 3 seconds")
    from app.agents.ingestion.docx_reader import DocxReader
    try:
        b = make_50page_docx()
        kb = len(b) / 1024
        print(f"           Fixture: {kb:.1f} KB")
        with patch(
            "app.agents.ingestion.docx_reader.storage_service.upload_raw",
            new_callable=AsyncMock, return_value="s3://k",
        ):
            t0 = time.perf_counter()
            ir = await DocxReader().read(b, "sc2_50page")
            elapsed = time.perf_counter() - t0
        ok = elapsed < 3.0
        record("SC-2   50-page .docx < 3 seconds", ok,
               f"elapsed={elapsed:.3f}s | elements={len(ir.elements)} | words={ir.word_count}")
    except Exception as exc:
        record("SC-2   50-page .docx < 3 seconds", False, str(exc))


# ── SC-3 ───────────────────────────────────────────────────────────────────────

async def sc3():
    print("\n[SC-3] Scanned PDF (image-only) triggers OcrReader")
    from app.agents.ingestion.pdf_reader import PdfReader
    from app.agents.ingestion.ocr_reader import OcrReader
    from app.schemas.ir import IRSchema

    ocr_called = False

    async def fake_ocr_read(self, file_bytes, job_id):
        nonlocal ocr_called
        ocr_called = True
        return IRSchema(
            source_format="pdf", elements=[], word_count=0,
            metadata={"ocr_processed": True, "page_count": 1},
        )

    try:
        # Fake fitz doc with zero-text page -> text_density = 0 < 0.5 threshold
        fake_doc = MagicMock()
        fake_doc.page_count = 1
        fake_page = MagicMock()
        fake_page.rect.width = 595.0
        fake_page.get_text.return_value = ""
        fake_doc.__iter__ = MagicMock(return_value=iter([fake_page]))
        fake_doc.close = MagicMock()

        with (
            patch("app.agents.ingestion.pdf_reader.fitz.open", return_value=fake_doc),
            patch.object(OcrReader, "read", fake_ocr_read),
        ):
            reader = PdfReader()
            ir = await reader.read(b"%PDF-dummy", "sc3_scanned")

        ok = ocr_called and ir.metadata.get("ocr_processed") is True
        record("SC-3   Scanned PDF triggers OcrReader", ok,
               f"ocr_called={ocr_called}, ocr_processed={ir.metadata.get('ocr_processed')}")
    except Exception as exc:
        record("SC-3   Scanned PDF triggers OcrReader", False, str(exc))


# ── SC-4 ───────────────────────────────────────────────────────────────────────

async def sc4():
    print("\n[SC-4] Extracted character count >= 95% of source document")
    from app.agents.ingestion.docx_reader import DocxReader
    from app.agents.ingestion.pdf_reader import PdfReader
    from app.agents.ingestion.text_reader import TextReader
    import fitz

    def alpha_count(text):
        return sum(1 for c in text if c.isalpha())

    # .txt
    try:
        content = (
            "# Research Paper\n\n"
            + "The quick brown fox jumps over the lazy dog. " * 200
            + "\n\n## References\n\nDoe (2022). Some Title. Journal.\n"
        ).encode("utf-8")
        src_a = alpha_count(content.decode())
        ir = await TextReader().read(content, "sc4_txt")
        ext = " ".join(e.raw_text for e in ir.elements)
        ext_a = alpha_count(ext)
        cov = ext_a / src_a if src_a else 0
        ok = cov >= 0.95
        record("SC-4a  .txt  coverage >= 95%", ok,
               f"{cov:.1%} covered ({ext_a}/{src_a} alpha chars)")
    except Exception as exc:
        record("SC-4a  .txt  coverage >= 95%", False, str(exc))

    # .docx
    try:
        from docx import Document
        doc = Document()
        sample = "Automated academic manuscript reformatting using agentic AI pipelines. " * 50
        for i in range(24):
            if i % 6 == 0:
                doc.add_heading(f"Chapter {i // 6 + 1}", level=1)
            else:
                doc.add_paragraph(f"Para {i}: " + sample)
        buf = BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        # Ground-truth: all para text inside the docx
        doc2 = Document(BytesIO(docx_bytes))
        src_text = " ".join(p.text for p in doc2.paragraphs)
        src_a = alpha_count(src_text)

        with patch(
            "app.agents.ingestion.docx_reader.storage_service.upload_raw",
            new_callable=AsyncMock, return_value="s3://k",
        ):
            ir = await DocxReader().read(docx_bytes, "sc4_docx")

        ext = " ".join(e.raw_text for e in ir.elements)
        ext_a = alpha_count(ext)
        cov = ext_a / src_a if src_a else 0
        ok = cov >= 0.95
        record("SC-4b  .docx coverage >= 95%", ok,
               f"{cov:.1%} covered ({ext_a}/{src_a} alpha chars)")
    except Exception as exc:
        record("SC-4b  .docx coverage >= 95%", False, str(exc))

    # .pdf
    try:
        pdf_bytes = make_pdf()
        raw_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        src_text = "".join(page.get_text() for page in raw_doc)
        raw_doc.close()
        src_a = alpha_count(src_text)

        with (
            patch("app.agents.ingestion.pdf_reader.storage_service.upload_raw",
                  new_callable=AsyncMock, return_value="s3://k"),
            patch.object(PdfReader, "_is_scanned", return_value=False),
        ):
            ir = await PdfReader().read(pdf_bytes, "sc4_pdf")

        ext = " ".join(e.raw_text for e in ir.elements)
        ext_a = alpha_count(ext)
        cov = ext_a / src_a if src_a else 0
        ok = cov >= 0.95
        record("SC-4c  .pdf  coverage >= 95%", ok,
               f"{cov:.1%} covered ({ext_a}/{src_a} alpha chars)")
    except Exception as exc:
        record("SC-4c  .pdf  coverage >= 95%", False, str(exc))


# ── SC-5 ───────────────────────────────────────────────────────────────────────

def sc5():
    print("\n[SC-5] pytest app/agents/ingestion/tests/ passes with 0 failures")
    res = subprocess.run(
        [sys.executable, "-m", "pytest",
         "app/agents/ingestion/tests/",
         "--asyncio-mode=auto", "-q", "--tb=short"],
        capture_output=True, text=True,
    )
    out = res.stdout + res.stderr
    lines = [l.strip() for l in out.splitlines() if l.strip()]
    summary = next(
        (l for l in reversed(lines) if "passed" in l or "failed" in l or "error" in l),
        out[-200:],
    )
    ok = res.returncode == 0 and "failed" not in summary and "error" not in summary.lower()
    record("SC-5   pytest 0 failures", ok, summary)
    if not ok:
        for line in lines[-25:]:
            print(f"  | {line}")


# ── Main ───────────────────────────────────────────────────────────────────────

async def main():
    print("=" * 65)
    print("  Agent Paperpal - DocIngestionAgent Success Criteria Report")
    print("=" * 65)

    await sc1()
    await sc2()
    await sc3()
    await sc4()
    sc5()

    total = len(results)
    passed = sum(1 for _, ok, _ in results if ok)
    failed = total - passed

    print("\n" + "=" * 65)
    print(f"  SUMMARY: {passed}/{total} criteria PASSED")
    print("=" * 65)
    for name, ok, detail in results:
        icon = "PASS" if ok else "FAIL"
        print(f"  [{icon}]  {name}")
        if detail:
            print(f"           {detail}")

    if failed == 0:
        print("\n  ALL SUCCESS CRITERIA MET!")
    else:
        print(f"\n  {failed} CRITERION/CRITERIA FAILED")
        sys.exit(1)


if __name__ == "__main__":
    asyncio.run(main())
