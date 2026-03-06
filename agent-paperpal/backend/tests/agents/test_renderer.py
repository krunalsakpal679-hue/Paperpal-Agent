# backend/tests/services/renderer/test_renderer.py
import pytest
from unittest.mock import AsyncMock, patch, MagicMock
from io import BytesIO
from docx import Document

from app.schemas.job_state import JobState, ChangeEntry
from app.schemas.ir import IRSchema, IRElement, ElementType, TextRun
from app.schemas.jro_schema import JROSchema, LayoutRules, HeadingRules, FigureRules, TableRules, AbstractRules, SectionRequirements

from app.services.renderer.style_applicator import StyleApplicator
from app.services.renderer.section_renderer import SectionRenderer
from app.services.renderer.table_renderer import TableRenderer
from app.services.renderer.figure_renderer import FigureRenderer
from app.services.renderer.renderer_service import RendererService

@pytest.fixture
def sample_jro():
    return JROSchema(
        journal_name="Test Journal",
        extraction_source="test",
        extraction_confidence=1.0,
        layout_rules=LayoutRules(
            font_name="Arial",
            font_size=11,
            line_spacing=1.5,
            margins={"top": 1.5, "bottom": 1.5, "left": 1.0, "right": 1.0}
        ),
        heading_rules=HeadingRules(),
        figure_rules=FigureRules(caption_position="bottom"),
        table_rules=TableRules(caption_position="top", border_style="simple"),
        abstract_rules=AbstractRules(),
        section_requirements=SectionRequirements()
    )

@pytest.fixture
def minimal_ir():
    return IRSchema(
        job_id="test_job",
        source_format="docx",
        elements=[
            IRElement(
                element_id="h1",
                element_type=ElementType.TITLE,
                raw_text="Test Document Title",
                content=[TextRun(text="Test Document Title")]
            ),
            IRElement(
                element_id="p1",
                element_type=ElementType.PARAGRAPH,
                raw_text="This is a body paragraph.",
                content=[TextRun(text="This is a body paragraph.")]
            ),
            IRElement(
                element_id="t1",
                element_type=ElementType.TABLE,
                raw_text="Table 1",
                metadata={"caption": "Table 1: Mock Table", "rows": [["1", "2"], ["3", "4"]]}
            )
        ]
    )

def test_style_applicator(sample_jro):
    applicator = StyleApplicator()
    doc = applicator.build_document_styles(sample_jro)
    
    # Check margins (1.5 inches = 2160000 EMUs, or access inch value)
    section = doc.sections[0]
    assert abs(section.top_margin.inches - 1.5) < 0.01
    
    # Check Normal style font
    normal_font = doc.styles['Normal'].font
    assert normal_font.name == "Arial"
    assert normal_font.size.pt == 11.0

def test_section_renderer():
    doc = Document()
    renderer = SectionRenderer()
    
    para_el = IRElement(element_id="1", element_type=ElementType.PARAGRAPH, raw_text="Testing paragraph")
    renderer.render(doc, para_el)
    
    title_el = IRElement(element_id="2", element_type=ElementType.TITLE, raw_text="Main Title")
    renderer.render(doc, title_el)
    
    assert len(doc.paragraphs) == 2
    assert "Testing paragraph" in doc.paragraphs[0].text
    assert "Main Title" in doc.paragraphs[1].text
    assert doc.paragraphs[1].style.name == 'Title' # Using level=0 creates Title style

def test_table_renderer(sample_jro):
    doc = Document()
    renderer = TableRenderer()
    
    table_el = IRElement(
        element_id="1",
        element_type=ElementType.TABLE,
        raw_text="",
        metadata={"caption": "Sample Table", "rows": [["A", "B"], ["C", "D"]]}
    )
    
    renderer.render(doc, table_el, sample_jro)
    
    # Needs to have the caption paragraph, then the table
    assert "Sample Table" in doc.paragraphs[-1].text
    assert doc.paragraphs[-1].style.name == 'Caption'
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 2
    assert doc.tables[0].cell(0, 0).text == "A"

@pytest.mark.asyncio
async def test_figure_renderer(sample_jro):
    doc = Document()
    mock_storage = MagicMock()
    # 1x1 white transparent PNG
    mock_storage.download_raw = AsyncMock(return_value=b'\\x89PNG\\r\\n\\x1a\\n\\x00\\x00\\x00\\rIHDR\\x00\\x00\\x00\\x01\\x00\\x00\\x00\\x01\\x08\\x06\\x00\\x00\\x00\\x1f\\x15\\xc4\\x89\\x00\\x00\\x00\\nIDATx\\x9cc\\x00\\x01\\x00\\x00\\x05\\x00\\x01\\r\\n-\\xb4\\x00\\x00\\x00\\x00IEND\\xaeB`\\x82')
    
    renderer = FigureRenderer(mock_storage)
    
    fig_el = IRElement(
        element_id="1",
        element_type=ElementType.FIGURE,
        raw_text="",
        metadata={"caption": "Sample Figure", "s3_key": "raw/img.png"}
    )
    
    await renderer.render(doc, fig_el, sample_jro)
    
    # Should have a paragraph for the image, then a paragraph for the caption (bottom pos)
    assert len(doc.paragraphs) >= 2
    assert "Sample Figure" in doc.paragraphs[-1].text

@pytest.mark.asyncio
async def test_renderer_service_integration(minimal_ir, sample_jro):
    # Setup state
    state = JobState(job_id="test_render_job", transformed_ir=minimal_ir, jro=sample_jro)
    state.change_log = [ChangeEntry(element_id="h1", field="text", old_value="", new_value="Test Document Title", rule_source="test")]
    
    service = RendererService()
    
    # Mock Storage
    with patch("app.services.renderer.renderer_service.storage_service.upload_output", new_callable=AsyncMock) as mock_upload, \
         patch("app.services.renderer.renderer_service.storage_service.get_signed_url", new_callable=AsyncMock) as mock_url:
        
        mock_upload.side_effect = ["key_docx", "key_json"]
        mock_url.side_effect = ["http://s3/test.docx", "http://s3/test.json"]
        
        urls = await service.render(state)
        
        assert "docx_url" in urls
        assert urls["docx_url"] == "http://s3/test.docx"
        assert "change_log_url" in urls
        
        # Verify call counts
        assert mock_upload.call_count == 2
        
        # Pull the docx bytes from the first upload call and verify it's a valid docx
        docx_bytes = mock_upload.mock_calls[0].args[2]
        doc = Document(BytesIO(docx_bytes))
        
        assert len(doc.paragraphs) >= 3
        # First header should be Title
        assert "Test Document Title" in doc.paragraphs[0].text
