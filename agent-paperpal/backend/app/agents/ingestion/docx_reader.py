# backend/app/agents/ingestion/docx_reader.py
"""
DocxReader — DOCX manuscript reader.

Extracts a structured Intermediate Representation from a .docx file using
python-docx. Handles paragraphs, headings, tables, inline images (uploaded
to S3), footnotes, and inline citations.
"""

import asyncio
import logging
import re
from io import BytesIO

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from app.schemas.ir import ElementType, IRElement, IRSchema, TextRun
from app.services.storage_service import storage_service

logger = logging.getLogger(__name__)

# ── Citation patterns ──────────────────────────────────────────────────────────
_AUTHOR_YEAR_PATTERN = re.compile(r"\([A-Z][a-z]+.*?\d{4}\)")
_NUMERIC_PATTERN = re.compile(r"\[\d+(?:,\s*\d+)*\]")

# ── Heading style prefixes (python-docx style names) ─────────────────────────
_HEADING_PREFIXES = ("Heading", "Title", "Subtitle")


class DocxReader:
    """
    Reads a .docx file and converts it to an IRSchema.

    Processing pipeline:
    - Core properties (title, author) → IRSchema.metadata
    - Paragraphs and headings       → IRElement list
    - Inline citations extracted    → metadata on each element
    - Tables                        → IRElement(TABLE)
    - Inline images                 → uploaded to S3, IRElement(FIGURE)
    - Footnotes                     → IRElement(FOOTNOTE)
    """

    async def read(self, file_bytes: bytes, job_id: str) -> IRSchema:
        """
        Parse a .docx document into an IRSchema.

        Args:
            file_bytes: Raw .docx bytes.
            job_id: UUID4 job identifier used for S3 upload paths.

        Returns:
            IRSchema populated with all extracted document elements.
        """
        doc = DocxDocument(BytesIO(file_bytes))
        elements: list[IRElement] = []
        word_count = 0

        # ── Core properties ────────────────────────────────────────────────────
        title, author = self._extract_core_properties(doc)

        # ── Paragraphs and headings ────────────────────────────────────────────
        for index, para in enumerate(doc.paragraphs):
            raw_text = para.text.strip()
            if not raw_text:
                continue

            element_type, level = self._classify_paragraph(para)
            runs = self._extract_runs(para)
            citations = self._extract_inline_citations(raw_text)
            word_count += len(raw_text.split())

            element = IRElement(
                element_id=f"p_{index}",
                element_type=element_type,
                content=runs,
                raw_text=raw_text,
                level=level,
                metadata={
                    "style_name": para.style.name,
                    "inline_citations": citations,
                },
            )
            elements.append(element)

        # ── Tables ─────────────────────────────────────────────────────────────
        for t_idx, table in enumerate(doc.tables):
            table_element = self._extract_table(table, t_idx)
            elements.append(table_element)

        # ── Inline images ──────────────────────────────────────────────────────
        figure_elements = await self._extract_images(doc, job_id)
        elements.extend(figure_elements)

        # ── Footnotes ──────────────────────────────────────────────────────────
        footnote_elements = self._extract_footnotes(doc)
        elements.extend(footnote_elements)

        logger.info(
            "DocxReader extracted %d elements, ~%d words for job %s",
            len(elements),
            word_count,
            job_id,
        )

        return IRSchema(
            document_title=title or "",
            authors=[author] if author else [],
            elements=elements,
            source_format="docx",
            word_count=word_count,
            metadata={
                "core_title": title,
                "core_author": author,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            },
        )

    # ── Private helpers ────────────────────────────────────────────────────────

    def _extract_core_properties(
        self, doc: DocxDocument
    ) -> tuple[str | None, str | None]:
        """Extract title and author from .docx core_properties."""
        props = doc.core_properties
        return getattr(props, "title", None), getattr(props, "author", None)

    def _classify_paragraph(
        self, para
    ) -> tuple[ElementType, int]:
        """
        Classify a paragraph as HEADING, TITLE, ABSTRACT, PARAGRAPH, etc.

        Returns:
            Tuple of (ElementType, heading_level).
            heading_level is 0 for non-headings.
        """
        style_name: str = para.style.name

        if style_name == "Title":
            return ElementType.TITLE, 1
        if style_name == "Abstract":
            return ElementType.ABSTRACT, 0

        # "Heading 1", "Heading 2", … "Heading 9"
        for prefix in _HEADING_PREFIXES:
            if style_name.startswith(prefix) and style_name != "Title":
                try:
                    level = int(style_name.split()[-1])
                except ValueError:
                    level = 1
                return ElementType.HEADING, level

        # Heuristic: bold-only short line → treat as heading
        raw_text = para.text.strip()
        runs = para.runs
        if (
            runs
            and all(r.bold for r in runs if r.text.strip())
            and len(raw_text) < 100
        ):
            return ElementType.HEADING, 2

        return ElementType.PARAGRAPH, 0

    def _extract_runs(self, para) -> list[TextRun]:
        """Convert python-docx runs into TextRun objects."""
        result: list[TextRun] = []
        for run in para.runs:
            if not run.text:
                continue
            font = run.font
            font_size_pt: float | None = None
            if font.size is not None:
                # font.size is in EMUs (914400 per inch); 12700 EMU per pt
                font_size_pt = font.size / 12700

            result.append(
                TextRun(
                    text=run.text,
                    bold=bool(run.bold),
                    italic=bool(run.italic),
                    underline=bool(run.underline),
                    superscript=bool(font.superscript),
                    subscript=bool(font.subscript),
                    font_name=font.name,
                    font_size_pt=font_size_pt,
                )
            )
        return result

    def _extract_inline_citations(self, text: str) -> list[str]:
        """Return list of citation strings found in the paragraph text."""
        citations: list[str] = []
        citations.extend(_AUTHOR_YEAR_PATTERN.findall(text))
        citations.extend(_NUMERIC_PATTERN.findall(text))
        return citations

    def _extract_table(self, table, t_idx: int) -> IRElement:
        """Convert a python-docx table into a TABLE IRElement."""
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])

        return IRElement(
            element_id=f"tbl_{t_idx}",
            element_type=ElementType.TABLE,
            content=[],
            raw_text="",
            metadata={"rows": rows, "row_count": len(rows)},
        )

    async def _extract_images(
        self, doc: DocxDocument, job_id: str
    ) -> list[IRElement]:
        """
        Extract inline images from the .docx package and upload them to S3.

        Images are found in the document's relationships and uploaded as
        jobs/{job_id}/raw_fig_{i}.png.
        """
        figure_elements: list[IRElement] = []
        image_parts = [
            rel.target_part
            for rel in doc.part.rels.values()
            if "image" in rel.reltype
        ]

        upload_tasks = []
        for i, img_part in enumerate(image_parts):
            img_bytes = img_part.blob
            filename = f"fig_{i}.png"
            upload_tasks.append(storage_service.upload_raw(job_id, filename, img_bytes))

        if upload_tasks:
            s3_keys = await asyncio.gather(*upload_tasks, return_exceptions=True)
            for i, result in enumerate(s3_keys):
                if isinstance(result, Exception):
                    logger.warning(
                        "Failed to upload image fig_%d for job %s: %s", i, job_id, result
                    )
                    s3_key = None
                else:
                    s3_key = result  # type: ignore[assignment]

                figure_elements.append(
                    IRElement(
                        element_id=f"fig_{i}",
                        element_type=ElementType.FIGURE,
                        content=[],
                        raw_text="",
                        metadata={"s3_key": s3_key, "index": i},
                    )
                )

        return figure_elements

    def _extract_footnotes(self, doc: DocxDocument) -> list[IRElement]:
        """
        Extract footnotes from the .docx XML part.

        Uses the internal footnotes part (word/footnotes.xml) if present.
        """
        footnote_elements: list[IRElement] = []
        try:
            footnotes_part = doc.part.footnotes  # type: ignore[attr-defined]
            footnote_elm = footnotes_part._element  # type: ignore[attr-defined]
            for fn_idx, fn in enumerate(footnote_elm.findall(qn("w:footnote"))):
                text_parts: list[str] = [
                    node.text or ""
                    for node in fn.iter(qn("w:t"))
                ]
                fn_text = "".join(text_parts).strip()
                fn_id = fn.get(qn("w:id"), str(fn_idx))
                if fn_text:
                    footnote_elements.append(
                        IRElement(
                            element_id=f"fn_{fn_id}",
                            element_type=ElementType.FOOTNOTE,
                            content=[TextRun(text=fn_text)],
                            raw_text=fn_text,
                        )
                    )
        except AttributeError:
            # No footnotes.xml part present — normal for many documents
            pass
        except Exception as exc:
            logger.debug("Footnote extraction skipped: %s", exc)

        return footnote_elements
