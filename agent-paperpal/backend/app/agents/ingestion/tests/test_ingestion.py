# backend/app/agents/ingestion/tests/test_ingestion.py
"""
Test suite — DocIngestionAgent and all sub-readers.

Tests cover:
  - FileValidator: size rejection, MIME rejection, corrupt file rejection
  - DocxReader: paragraphs, headings, tables, inline citations
  - PdfReader: text extraction, heading inference
  - OcrReader: scanned PDF detection trigger (mocked)
  - TextReader: Markdown headings, ALL CAPS, colon-label, numbered, reference section
  - DocIngestionAgent.run(): full pipeline with mocked S3 + Redis

All external I/O (S3, Redis, langdetect, pdf2image, pytesseract) is mocked.
"""

from __future__ import annotations

import io
from io import BytesIO
from typing import Any
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
import pytest_asyncio

# ──────────────────────────────────────────────────────────────────────────────
# Fixtures — document factories
# ──────────────────────────────────────────────────────────────────────────────


@pytest.fixture()
def minimal_docx_bytes() -> bytes:
    """Create a minimal in-memory .docx with heading + body + table."""
    from docx import Document as DocxDocument
    from docx.shared import Pt

    doc = DocxDocument()
    doc.core_properties.title = "Test Paper"
    doc.core_properties.author = "Jane Doe"

    # Title paragraph (Heading 1 style — word uses "Heading 1")
    heading = doc.add_heading("Introduction", level=1)
    _ = heading  # captured but not used further

    # Body paragraph with author-year citation
    para = doc.add_paragraph(
        "This work builds on prior research (Smith et al. 2023) extensively."
    )
    run = para.add_run()  # extra run
    run.bold = False
    run.font.size = Pt(12)

    # Italic paragraph
    italic_para = doc.add_paragraph("")
    italic_run = italic_para.add_run("Emphasised claim here.")
    italic_run.italic = True

    # A small table
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Col A"
    table.cell(0, 1).text = "Col B"
    table.cell(0, 2).text = "Col C"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    table.cell(1, 2).text = "3"

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture()
def minimal_pdf_bytes() -> bytes:
    """Create a minimal 2-page text PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.pdfgen import canvas
    except ImportError:
        pytest.skip("reportlab not installed")

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)

    # Page 1
    c.setFont("Helvetica-Bold", 18)
    c.drawString(72, 700, "Abstract")
    c.setFont("Helvetica", 12)
    c.drawString(
        72, 670, "This paper investigates methods for academic reformatting."
    )
    c.drawString(72, 650, "We propose a novel pipeline using LLM-based agents.")
    c.showPage()

    # Page 2
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 700, "Methods")
    c.setFont("Helvetica", 12)
    c.drawString(72, 670, "We collected 1000 manuscripts from open-access repositories.")
    c.showPage()

    c.save()
    return buf.getvalue()


@pytest.fixture()
def oversized_file_bytes() -> bytes:
    """Generate a byte string that exceeds a 1 MB fictitious limit."""
    # 2 MB of zeros
    return b"\x00" * (2 * 1024 * 1024)


@pytest.fixture()
def markdown_text_bytes() -> bytes:
    """A simple Markdown document as bytes."""
    content = """\
# My Research Paper

## Abstract

This paper presents a comprehensive study of NLP pipelines (Jones 2022).

## Introduction

Natural language processing has advanced rapidly in recent years [1][2].

### Background

Early work by Smith et al. (2019) established the baseline.

## References

Jones, A. (2022). NLP in the wild. *Journal of AI*. 10(2), 100–120.
Smith, B., & Lee, C. (2019). Baseline methods. *Proc. ACL*.
"""
    return content.encode("utf-8")


@pytest.fixture()
def plain_text_bytes() -> bytes:
    """A plain-text document with ALL CAPS headings."""
    content = """\
INTRODUCTION

This section introduces the topic.
Background material is provided.

METHODS

We used Python 3.12 for all experiments.
The pipeline consists of five stages.

REFERENCES

Smith (2020). Some paper. Journal X.
"""
    return content.encode("utf-8")


# ──────────────────────────────────────────────────────────────────────────────
# FileValidator tests
# ──────────────────────────────────────────────────────────────────────────────


class TestFileValidator:
    """Unit tests for FileValidator."""

    def _make_validator(self, max_mb: int = 50):
        from app.agents.ingestion.file_validator import FileValidator

        with patch("app.agents.ingestion.file_validator.settings") as mock_settings:
            mock_settings.MAX_FILE_SIZE_MB = max_mb
            mock_settings.max_file_size_bytes = max_mb * 1024 * 1024
            validator = FileValidator()
        return validator

    def test_accepts_valid_docx(self, minimal_docx_bytes: bytes) -> None:
        from app.agents.ingestion.file_validator import FileValidator

        validator = FileValidator()
        with (
            patch("app.agents.ingestion.file_validator.settings") as ms,
            patch("app.agents.ingestion.file_validator.magic.from_buffer") as mock_magic,
        ):
            ms.MAX_FILE_SIZE_MB = 50
            ms.max_file_size_bytes = 50 * 1024 * 1024
            mock_magic.return_value = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            errors = validator.validate("test.docx", minimal_docx_bytes)

        assert errors == [], f"Expected no errors, got: {errors}"

    def test_rejects_oversized_file(self, oversized_file_bytes: bytes) -> None:
        from app.agents.ingestion.file_validator import FileValidator

        validator = FileValidator()
        with patch("app.agents.ingestion.file_validator.settings") as ms:
            ms.MAX_FILE_SIZE_MB = 1
            ms.max_file_size_bytes = 1 * 1024 * 1024  # 1 MB limit
            errors = validator.validate("big_paper.docx", oversized_file_bytes)

        assert len(errors) == 1
        assert errors[0]["code"] == "FILE_TOO_LARGE"

    def test_rejects_unsupported_mime(self) -> None:
        from app.agents.ingestion.file_validator import FileValidator

        validator = FileValidator()
        fake_exe_bytes = b"MZ\x90\x00" + b"\x00" * 100  # .exe magic bytes

        with (
            patch("app.agents.ingestion.file_validator.settings") as ms,
            patch("app.agents.ingestion.file_validator.magic.from_buffer") as mock_magic,
        ):
            ms.MAX_FILE_SIZE_MB = 50
            ms.max_file_size_bytes = 50 * 1024 * 1024
            mock_magic.return_value = "application/x-msdownload"
            errors = validator.validate("malware.exe", fake_exe_bytes)

        assert len(errors) == 1
        assert errors[0]["code"] == "UNSUPPORTED_FILE_TYPE"

    def test_rejects_corrupt_pdf(self) -> None:
        from app.agents.ingestion.file_validator import FileValidator

        validator = FileValidator()
        corrupt_bytes = b"%PDF-1.4 corrupt nonsense!!"

        with (
            patch("app.agents.ingestion.file_validator.settings") as ms,
            patch("app.agents.ingestion.file_validator.magic.from_buffer") as mock_magic,
        ):
            ms.MAX_FILE_SIZE_MB = 50
            ms.max_file_size_bytes = 50 * 1024 * 1024
            mock_magic.return_value = "application/pdf"
            errors = validator.validate("corrupt.pdf", corrupt_bytes)

        assert len(errors) >= 1
        assert errors[0]["code"] in ("CORRUPT_PDF", "PDF_PARSE_ERROR")

    def test_rejects_corrupt_docx(self) -> None:
        from app.agents.ingestion.file_validator import FileValidator

        validator = FileValidator()
        not_a_docx = b"This is definitely not a docx file at all."

        with (
            patch("app.agents.ingestion.file_validator.settings") as ms,
            patch("app.agents.ingestion.file_validator.magic.from_buffer") as mock_magic,
        ):
            ms.MAX_FILE_SIZE_MB = 50
            ms.max_file_size_bytes = 50 * 1024 * 1024
            mock_magic.return_value = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            errors = validator.validate("fake.docx", not_a_docx)

        assert len(errors) >= 1
        assert errors[0]["code"] in ("CORRUPT_DOCX", "DOCX_PARSE_ERROR")


# ──────────────────────────────────────────────────────────────────────────────
# DocxReader tests
# ──────────────────────────────────────────────────────────────────────────────


class TestDocxReader:
    """Unit tests for DocxReader."""

    @pytest.mark.asyncio
    async def test_extracts_paragraphs_and_heading(
        self, minimal_docx_bytes: bytes
    ) -> None:
        from app.agents.ingestion.docx_reader import DocxReader
        from app.schemas.ir import ElementType

        with patch(
            "app.agents.ingestion.docx_reader.storage_service.upload_raw",
            new_callable=AsyncMock,
            return_value="jobs/test_job/raw_fig_0.png",
        ):
            reader = DocxReader()
            ir = await reader.read(minimal_docx_bytes, "test_job")

        assert ir.source_format == "docx"
        assert len(ir.elements) > 0
        # Must contain at least one HEADING element
        headings = [e for e in ir.elements if e.element_type == ElementType.HEADING]
        assert len(headings) >= 1
        assert "Introduction" in headings[0].raw_text

    @pytest.mark.asyncio
    async def test_extracts_inline_citations(self, minimal_docx_bytes: bytes) -> None:
        from app.agents.ingestion.docx_reader import DocxReader

        with patch(
            "app.agents.ingestion.docx_reader.storage_service.upload_raw",
            new_callable=AsyncMock,
            return_value="s3key",
        ):
            reader = DocxReader()
            ir = await reader.read(minimal_docx_bytes, "job_cite")

        # Find the paragraph containing the citation text
        citation_paras = [
            e
            for e in ir.elements
            if "Smith" in e.raw_text and e.metadata.get("inline_citations")
        ]
        assert len(citation_paras) >= 1

    @pytest.mark.asyncio
    async def test_extracts_table(self, minimal_docx_bytes: bytes) -> None:
        from app.agents.ingestion.docx_reader import DocxReader
        from app.schemas.ir import ElementType

        with patch(
            "app.agents.ingestion.docx_reader.storage_service.upload_raw",
            new_callable=AsyncMock,
            return_value="s3key",
        ):
            reader = DocxReader()
            ir = await reader.read(minimal_docx_bytes, "job_table")

        table_elements = [
            e for e in ir.elements if e.element_type == ElementType.TABLE
        ]
        assert len(table_elements) >= 1
        rows = table_elements[0].metadata["rows"]
        assert rows[0] == ["Col A", "Col B", "Col C"]

    @pytest.mark.asyncio
    async def test_ir_validates_with_pydantic(self, minimal_docx_bytes: bytes) -> None:
        from app.agents.ingestion.docx_reader import DocxReader
        from app.schemas.ir import IRSchema

        with patch(
            "app.agents.ingestion.docx_reader.storage_service.upload_raw",
            new_callable=AsyncMock,
            return_value="s3key",
        ):
            reader = DocxReader()
            ir = await reader.read(minimal_docx_bytes, "job_validate")

        # Full round-trip Pydantic validation
        ir_dict = ir.model_dump()
        revalidated = IRSchema.model_validate(ir_dict)
        assert revalidated.source_format == "docx"

    @pytest.mark.asyncio
    async def test_word_count_nonzero(self, minimal_docx_bytes: bytes) -> None:
        from app.agents.ingestion.docx_reader import DocxReader

        with patch(
            "app.agents.ingestion.docx_reader.storage_service.upload_raw",
            new_callable=AsyncMock,
            return_value="s3key",
        ):
            reader = DocxReader()
            ir = await reader.read(minimal_docx_bytes, "job_wc")

        assert ir.word_count > 0

    @pytest.mark.asyncio
    async def test_title_from_core_properties(self, minimal_docx_bytes: bytes) -> None:
        from app.agents.ingestion.docx_reader import DocxReader

        with patch(
            "app.agents.ingestion.docx_reader.storage_service.upload_raw",
            new_callable=AsyncMock,
            return_value="s3key",
        ):
            reader = DocxReader()
            ir = await reader.read(minimal_docx_bytes, "job_title")

        assert ir.document_title == "Test Paper"
        assert "Jane Doe" in ir.authors


# ──────────────────────────────────────────────────────────────────────────────
# PdfReader tests
# ──────────────────────────────────────────────────────────────────────────────


class TestPdfReader:
    """Unit tests for PdfReader.

    NOTE: All PDF tests use fitz (PyMuPDF) directly which is installed.
    Tests that would trigger OcrReader (pdf2image/poppler) fully mock the
    OcrReader.read path to avoid the missing Poppler system dependency on Windows.
    """

    @pytest.mark.asyncio
    async def test_extracts_text_elements(self, minimal_pdf_bytes: bytes) -> None:
        from app.agents.ingestion.pdf_reader import PdfReader

        with (
            patch(
                "app.agents.ingestion.pdf_reader.storage_service.upload_raw",
                new_callable=AsyncMock,
                return_value="s3key",
            ),
            # Ensure PdfReader does NOT think this is a scanned PDF
            # (our reportlab fixture has real text, but add guard for safety)
            patch.object(PdfReader, "_is_scanned", return_value=False),
        ):
            reader = PdfReader()
            ir = await reader.read(minimal_pdf_bytes, "job_pdf")

        assert ir.source_format == "pdf"
        assert len(ir.elements) > 0
        all_text = " ".join(e.raw_text for e in ir.elements)
        assert len(all_text) > 10  # Some text was extracted

    @pytest.mark.asyncio
    async def test_page_count_in_metadata(self, minimal_pdf_bytes: bytes) -> None:
        from app.agents.ingestion.pdf_reader import PdfReader

        with (
            patch(
                "app.agents.ingestion.pdf_reader.storage_service.upload_raw",
                new_callable=AsyncMock,
                return_value="s3key",
            ),
            patch.object(PdfReader, "_is_scanned", return_value=False),
        ):
            reader = PdfReader()
            ir = await reader.read(minimal_pdf_bytes, "job_pages")

        assert ir.metadata.get("page_count") == 2

    @pytest.mark.asyncio
    async def test_scanned_pdf_triggers_ocr(self) -> None:
        """
        Image-only PDF (zero text density) must trigger OcrReader.
        We mock fitz.open + _is_scanned=True and OcrReader.read to avoid
        needing Poppler or a real PDF on Windows.
        """
        from app.agents.ingestion.pdf_reader import PdfReader
        from app.agents.ingestion.ocr_reader import OcrReader
        from app.schemas.ir import IRSchema

        mock_ir = IRSchema(
            source_format="pdf",
            elements=[],
            word_count=0,
            metadata={"ocr_processed": True},
        )

        # Build a fake fitz document so fitz.open doesn't choke on our dummy bytes
        fake_doc = MagicMock()
        fake_doc.page_count = 1
        fake_doc.__iter__ = MagicMock(return_value=iter([]))
        fake_doc.close = MagicMock()

        with (
            patch("app.agents.ingestion.pdf_reader.fitz.open", return_value=fake_doc),
            patch.object(PdfReader, "_is_scanned", return_value=True),
            patch.object(
                OcrReader,
                "read",
                new_callable=AsyncMock,
                return_value=mock_ir,
            ) as mock_ocr,
        ):
            reader = PdfReader()
            ir = await reader.read(b"%PDF-1.4 dummy", "job_ocr_trigger")

        mock_ocr.assert_called_once()
        assert ir.metadata.get("ocr_processed") is True


    @pytest.mark.asyncio
    async def test_ir_validates_pydantic(self, minimal_pdf_bytes: bytes) -> None:
        from app.agents.ingestion.pdf_reader import PdfReader
        from app.schemas.ir import IRSchema

        with (
            patch(
                "app.agents.ingestion.pdf_reader.storage_service.upload_raw",
                new_callable=AsyncMock,
                return_value="s3key",
            ),
            patch.object(PdfReader, "_is_scanned", return_value=False),
        ):
            reader = PdfReader()
            ir = await reader.read(minimal_pdf_bytes, "job_pdf_v")

        revalidated = IRSchema.model_validate(ir.model_dump())
        assert revalidated.source_format == "pdf"


# ──────────────────────────────────────────────────────────────────────────────
# TextReader tests
# ──────────────────────────────────────────────────────────────────────────────


class TestTextReader:
    """Unit tests for TextReader."""

    @pytest.mark.asyncio
    async def test_markdown_heading_detection(self, markdown_text_bytes: bytes) -> None:
        from app.agents.ingestion.text_reader import TextReader
        from app.schemas.ir import ElementType

        reader = TextReader()
        ir = await reader.read(markdown_text_bytes, "job_md")

        headings = [e for e in ir.elements if e.element_type == ElementType.HEADING]
        heading_texts = [h.raw_text for h in headings]

        assert any("My Research Paper" in t for t in heading_texts)
        assert any("Abstract" in t for t in heading_texts)
        assert any("Introduction" in t for t in heading_texts)

    @pytest.mark.asyncio
    async def test_all_caps_heading(self, plain_text_bytes: bytes) -> None:
        from app.agents.ingestion.text_reader import TextReader
        from app.schemas.ir import ElementType

        reader = TextReader()
        ir = await reader.read(plain_text_bytes, "job_plain")

        headings = [e for e in ir.elements if e.element_type == ElementType.HEADING]
        heading_texts = [h.raw_text for h in headings]

        assert any("INTRODUCTION" in t for t in heading_texts)
        assert any("METHODS" in t for t in heading_texts)

    @pytest.mark.asyncio
    async def test_reference_section_classification(
        self, markdown_text_bytes: bytes
    ) -> None:
        from app.agents.ingestion.text_reader import TextReader
        from app.schemas.ir import ElementType

        reader = TextReader()
        ir = await reader.read(markdown_text_bytes, "job_refs")

        reference_elements = [
            e
            for e in ir.elements
            if e.element_type == ElementType.REFERENCE
        ]
        assert len(reference_elements) >= 1
        all_ref_text = " ".join(e.raw_text for e in reference_elements)
        assert "Jones" in all_ref_text or "Smith" in all_ref_text

    @pytest.mark.asyncio
    async def test_utf8_decode(self) -> None:
        from app.agents.ingestion.text_reader import TextReader

        utf8_bytes = "# Héllo Wörld\n\nSome paragraph.".encode("utf-8")
        reader = TextReader()
        ir = await reader.read(utf8_bytes, "job_utf8")
        assert "Héllo Wörld" in ir.document_title or any(
            "Héllo" in e.raw_text for e in ir.elements
        )

    @pytest.mark.asyncio
    async def test_latin1_fallback(self) -> None:
        from app.agents.ingestion.text_reader import TextReader

        latin1_bytes = "Caf\xe9 and naïve words.\n\nAnother paragraph.".encode(
            "latin-1"
        )
        reader = TextReader()
        ir = await reader.read(latin1_bytes, "job_latin1")
        assert ir.word_count > 0

    @pytest.mark.asyncio
    async def test_numbered_heading(self) -> None:
        from app.agents.ingestion.text_reader import TextReader
        from app.schemas.ir import ElementType

        content = b"1. Introduction\n\nSome body text here.\n\n2.1 Background\n\nMore text."
        reader = TextReader()
        ir = await reader.read(content, "job_num")

        headings = [e for e in ir.elements if e.element_type == ElementType.HEADING]
        assert len(headings) >= 2

    @pytest.mark.asyncio
    async def test_ir_validates_pydantic(self, markdown_text_bytes: bytes) -> None:
        from app.agents.ingestion.text_reader import TextReader
        from app.schemas.ir import IRSchema

        reader = TextReader()
        ir = await reader.read(markdown_text_bytes, "job_txt_v")

        revalidated = IRSchema.model_validate(ir.model_dump())
        assert revalidated.source_format == "txt"

    @pytest.mark.asyncio
    async def test_character_coverage_threshold(self) -> None:
        """
        Verify ≥ 95% of source characters are captured in extracted text.
        We measure by comparing source char count to extracted raw_text sum.
        """
        from app.agents.ingestion.text_reader import TextReader

        body = "Lorem ipsum dolor sit amet. " * 200  # ~5600 chars of body
        content = f"# Title\n\n{body}\n\n## References\n\nRef 1.\n".encode("utf-8")
        source_alpha_count = sum(1 for c in content.decode() if c.isalpha())

        reader = TextReader()
        ir = await reader.read(content, "job_coverage")

        extracted = " ".join(e.raw_text for e in ir.elements)
        extracted_alpha_count = sum(1 for c in extracted if c.isalpha())

        coverage = extracted_alpha_count / source_alpha_count
        assert coverage >= 0.95, (
            f"Character coverage {coverage:.2%} is below 95% threshold"
        )


# ──────────────────────────────────────────────────────────────────────────────
# DocIngestionAgent integration tests
# ──────────────────────────────────────────────────────────────────────────────


class TestDocIngestionAgent:
    """Integration tests for the full DocIngestionAgent.run() pipeline."""

    def _make_state(
        self,
        raw_s3_key: str = "jobs/job123/raw_paper.docx",
        source_format: str = "docx",
    ):
        from app.schemas.job_state import JobState

        return JobState(
            job_id="job123",
            metadata={
                "raw_s3_key": raw_s3_key,
                "source_format": source_format,
            },
        )

    @pytest.mark.asyncio
    async def test_successful_docx_run(self, minimal_docx_bytes: bytes) -> None:
        from app.agents.ingestion.agent import DocIngestionAgent
        from app.schemas.job_state import JobStatus

        state = self._make_state(source_format="docx")

        with (
            patch(
                "app.agents.ingestion.agent.storage_service.download_raw",
                new_callable=AsyncMock,
                return_value=minimal_docx_bytes,
            ),
            patch(
                "app.agents.ingestion.docx_reader.storage_service.upload_raw",
                new_callable=AsyncMock,
                return_value="s3key",
            ),
            patch(
                "app.agents.ingestion.agent.cache_service.publish_progress",
                new_callable=AsyncMock,
            ),
            # Mock libmagic to return the correct DOCX MIME type
            patch("app.agents.ingestion.file_validator.magic.from_buffer") as mock_magic,
            # Mock FileValidator settings (size limit)
            patch("app.agents.ingestion.file_validator.settings") as fv_ms,
            # Mock langdetect
            patch("app.agents.ingestion.agent.detect", return_value="en"),
        ):
            mock_magic.return_value = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            fv_ms.MAX_FILE_SIZE_MB = 50
            fv_ms.max_file_size_bytes = 50 * 1024 * 1024

            agent = DocIngestionAgent()
            updated_state = await agent.run(state)

        assert updated_state.raw_ir is not None
        assert updated_state.status == JobStatus.INGESTING
        assert updated_state.progress_pct == 20.0
        assert updated_state.errors == []

    @pytest.mark.asyncio
    async def test_validation_failure_sets_failed_status(self) -> None:
        from app.agents.ingestion.agent import DocIngestionAgent
        from app.schemas.job_state import JobStatus

        state = self._make_state()

        with (
            patch(
                "app.agents.ingestion.agent.storage_service.download_raw",
                new_callable=AsyncMock,
                return_value=b"bad bytes",
            ),
            patch(
                "app.agents.ingestion.file_validator.magic.from_buffer",
                return_value="application/x-msdownload",
            ),
            patch("app.agents.ingestion.file_validator.settings") as fv_ms,
        ):
            fv_ms.MAX_FILE_SIZE_MB = 50
            fv_ms.max_file_size_bytes = 50 * 1024 * 1024

            agent = DocIngestionAgent()
            updated_state = await agent.run(state)

        assert updated_state.status == JobStatus.FAILED
        assert len(updated_state.errors) >= 1

    @pytest.mark.asyncio
    async def test_missing_s3_key_raises_error(self) -> None:
        from app.agents.ingestion.agent import DocIngestionAgent
        from app.schemas.job_state import JobState, JobStatus

        state = JobState(job_id="job_err", metadata={})

        with patch(
            "app.agents.ingestion.agent.storage_service.download_raw",
            new_callable=AsyncMock,
        ) as mock_dl:
            agent = DocIngestionAgent()
            updated_state = await agent.run(state)

        assert updated_state.status == JobStatus.FAILED
        assert any("raw_s3_key" in e.message for e in updated_state.errors)
        # download_raw should not be called if s3_key is missing
        mock_dl.assert_not_called()

    @pytest.mark.asyncio
    async def test_s3_download_failure_sets_failed_status(self) -> None:
        from app.agents.ingestion.agent import DocIngestionAgent
        from app.schemas.job_state import JobStatus

        state = self._make_state()

        with patch(
            "app.agents.ingestion.agent.storage_service.download_raw",
            new_callable=AsyncMock,
            side_effect=ConnectionError("S3 not reachable"),
        ):
            agent = DocIngestionAgent()
            updated_state = await agent.run(state)

        assert updated_state.status == JobStatus.FAILED
        assert any("ConnectionError" in e.error_type for e in updated_state.errors)

    @pytest.mark.asyncio
    async def test_progress_published_on_success(
        self, minimal_docx_bytes: bytes
    ) -> None:
        from app.agents.ingestion.agent import DocIngestionAgent

        state = self._make_state(source_format="docx")

        with (
            patch(
                "app.agents.ingestion.agent.storage_service.download_raw",
                new_callable=AsyncMock,
                return_value=minimal_docx_bytes,
            ),
            patch(
                "app.agents.ingestion.docx_reader.storage_service.upload_raw",
                new_callable=AsyncMock,
                return_value="s3key",
            ),
            patch(
                "app.agents.ingestion.agent.cache_service.publish_progress",
                new_callable=AsyncMock,
            ) as mock_publish,
            patch("app.agents.ingestion.file_validator.magic.from_buffer") as mock_magic,
            patch("app.agents.ingestion.file_validator.settings") as fv_ms,
            patch("app.agents.ingestion.agent.detect", return_value="en"),
        ):
            mock_magic.return_value = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            fv_ms.MAX_FILE_SIZE_MB = 50
            fv_ms.max_file_size_bytes = 50 * 1024 * 1024

            agent = DocIngestionAgent()
            await agent.run(state)

        mock_publish.assert_called_once()
        call_kwargs = mock_publish.call_args
        assert call_kwargs.kwargs["job_id"] == "job123"
        event = call_kwargs.kwargs["event_dict"]
        assert event["agent"] == "ingestion"
        assert event["pct"] == 20

    @pytest.mark.asyncio
    async def test_txt_routing(self, markdown_text_bytes: bytes) -> None:
        from app.agents.ingestion.agent import DocIngestionAgent
        from app.schemas.job_state import JobStatus

        state = self._make_state(
            raw_s3_key="jobs/job_txt/raw_paper.md", source_format="md"
        )

        with (
            patch(
                "app.agents.ingestion.agent.storage_service.download_raw",
                new_callable=AsyncMock,
                return_value=markdown_text_bytes,
            ),
            patch(
                "app.agents.ingestion.agent.cache_service.publish_progress",
                new_callable=AsyncMock,
            ),
            patch("app.agents.ingestion.file_validator.magic.from_buffer") as mock_magic,
            patch("app.agents.ingestion.file_validator.settings") as fv_ms,
            patch("app.agents.ingestion.agent.detect", return_value="en"),
        ):
            mock_magic.return_value = "text/markdown"
            fv_ms.MAX_FILE_SIZE_MB = 50
            fv_ms.max_file_size_bytes = 50 * 1024 * 1024

            agent = DocIngestionAgent()
            updated_state = await agent.run(state)

        assert updated_state.status == JobStatus.INGESTING
        assert updated_state.raw_ir is not None
        assert updated_state.raw_ir.source_format == "txt"
