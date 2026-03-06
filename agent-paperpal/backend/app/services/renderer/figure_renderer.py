# backend/app/services/renderer/figure_renderer.py
import logging
from io import BytesIO
from docx.document import Document
from docx.shared import Inches
from app.schemas.ir import IRElement
from app.schemas.jro_schema import JROSchema

class FigureRenderer:
    def __init__(self, storage_service):
         self.storage = storage_service
         self.logger = logging.getLogger(__name__)

    async def render(self, doc: Document, ir_figure: IRElement, jro: JROSchema) -> None:
        """Downloads a figure from S3 and adds it to the .docx."""

        caption = ir_figure.metadata.get("caption", "Figure")
        s3_key = ir_figure.metadata.get("s3_key")
        caption_pos = jro.figure_rules.caption_position if jro.figure_rules else "bottom"

        # 1. Caption Above
        if caption_pos in ("top", "above"):
             doc.add_paragraph(caption, style="Caption")

        # 2. Add Image
        if s3_key:
             try:
                 img_bytes = await self.storage.download_raw(s3_key)
                 img_stream = BytesIO(img_bytes)
                 doc.add_picture(img_stream, width=Inches(5.0))
                 # Center aligning the last paragraph (the picture)
                 last_paragraph = doc.paragraphs[-1]
                 last_paragraph.alignment = 1 # WD_ALIGN_PARAGRAPH.CENTER
             except Exception as e:
                 self.logger.error("Failed to append figure '%s' from S3: %s", caption, e)
                 doc.add_paragraph(f"[Image Placeholder for: {caption}]", style="Normal")
        
        # 3. Caption Below
        if caption_pos in ("bottom", "below", "any"):
             doc.add_paragraph(caption, style="Caption")

