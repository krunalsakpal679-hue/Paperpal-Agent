# backend/app/services/renderer/table_renderer.py
from docx.document import Document
from app.schemas.ir import IRElement
from app.schemas.jro_schema import JROSchema

class TableRenderer:
    def render(self, doc: Document, ir_table: IRElement, jro: JROSchema) -> None:
        """Render a table based on the JRO rules."""
        
        caption = ir_table.metadata.get("caption", "Table")
        rows = ir_table.metadata.get("rows", [[]])
        
        caption_pos = jro.table_rules.caption_position if jro.table_rules else "top"
        border_style = jro.table_rules.border_style if jro.table_rules else "simple"

        # 1. Caption Above
        if caption_pos in ("top", "above", "any"):
            doc.add_paragraph(caption, style="Caption")

        # 2. Add table
        max_cols = max((len(r) for r in rows), default=0)
        table = doc.add_table(rows=len(rows), cols=max_cols)
        
        # 3. Apply Border style
        if border_style == "simple" or border_style == "single":
            table.style = 'Table Grid'
            
        # 4. Fill cells
        for i, row in enumerate(rows):
            for j, text in enumerate(row):
                if j < len(table.columns): # bounds check
                    table.cell(i, j).text = str(text)

        # 5. Caption Below
        if caption_pos == "bottom":
             doc.add_paragraph(caption, style="Caption")
