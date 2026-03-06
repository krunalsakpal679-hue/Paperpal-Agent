# backend/app/services/renderer/section_renderer.py
from docx.document import Document
from app.schemas.ir import IRElement, ElementType

class SectionRenderer:
    def render(self, doc: Document, paragraph: IRElement) -> None:
        """Route paragraph rendering by element_type."""
        # Using abstract_rules or similar is not needed unless appending specifically
        text = paragraph.raw_text
        if not text:
            return

        if paragraph.element_type == ElementType.TITLE:
            doc.add_heading(text, level=0)
        
        elif paragraph.element_type == ElementType.ABSTRACT:
            doc.add_heading('Abstract', level=1)
            doc.add_paragraph(text, style='Normal')
        
        elif paragraph.element_type == ElementType.KEYWORD:
            doc.add_paragraph(f'Keywords: {text}', style='Normal')
        
        elif paragraph.element_type == ElementType.HEADING:
            level = paragraph.level if paragraph.level else 1
            doc.add_heading(text, level=min(level, 3)) # Max level 3 usually supported by default
        
        elif paragraph.element_type == ElementType.PARAGRAPH:
            para = doc.add_paragraph(text, style='Normal')
        
        elif paragraph.element_type == ElementType.FIGURE_CAPTION:
            doc.add_paragraph(text, style='Caption')
        
        elif paragraph.element_type == ElementType.TABLE_CAPTION:
            doc.add_paragraph(text, style='Caption')
        
        elif paragraph.element_type == ElementType.REFERENCE:
            doc.add_paragraph(text, style='Normal')
        
        elif paragraph.element_type == ElementType.FOOTNOTE:
            # Add to end of document (since python-docx doesn't easily support inline footnotes initially)
            doc.add_paragraph(f"Footnote {paragraph.metadata.get('number', '')}: {text}", style='Normal')
            
        else:
            # Default to text
            doc.add_paragraph(text, style='Normal')
