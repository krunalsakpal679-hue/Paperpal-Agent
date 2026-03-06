"""
validate_success_criteria.py
============================
Runs all 5 ingestion-agent success criteria and prints a PASS/FAIL report.

Criteria:
  SC-1  All 3 formats (.docx / .pdf / .txt|.md) produce IR that pass model_validate()
  SC-2  50-page .docx completes in < 3 seconds
  SC-3  Scanned PDF (image-only) triggers OcrReader delegation
  SC-4  Character count of extracted text >= 95% of source document
  SC-5  pytest app/agents/ingestion/tests/ passes with 0 failures

Usage:
    cd backend
    python validate_success_criteria.py
"""

import asyncio
import subprocess
import sys
import textwrap
import time
from io import BytesIO
from unittest.mock import AsyncMock, MagicMock, patch

# ── Console helpers ────────────────────────────────────────────────────────────
GREEN  = "\033[92m"
RED    = "\033[91m"
YELLOW = "\033[93m"
RESET  = "\033[0m"
BOLD   = "\033[1m"

results: list[tuple[str, bool, str]] = []   # (name, passed, detail)


def record(name: str, passed: bool, detail: str = "") -> None:
    results.append((name, passed, detail))
    icon = f"{GREEN}✅ PASS{RESET}" if passed else f"{RED}❌ FAIL{RESET}"
    print(f"  {icon}  {name}")
    if detail:
        for line in textwrap.wrap(detail, 90):
            print(f"         {line}")


# ── Fixtures ────────────────────────────────────────────────────────────────────

def make_docx(num_paragraphs: int = 5) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.core_properties.title = "Benchmark Paper"
    doc.core_properties.author = "Test Author"
    for i in range(num_paragraphs):
        if i % 8 == 0:
            doc.add_heading(f"Section {i // 8 + 1}", level=1)
        else:
            doc.add_paragraph(
                f"Paragraph {i}: " + ("Lorem ipsum dolor sit amet, consectetur "
                                      "adipiscing elit. Sed do eiusmod tempor. ") * 3
            )
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_50page_docx() -> bytes:
    """~50 pages: 400 body paragraphs + 50 headings."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.core_properties.title = "50-Page Benchmark"
    doc.core_properties.author = "Benchmark Author"
    # ~8 paragraphs per page → 400 body paras + 50 headings = ~56 pages
    for section in range(50):
        doc.add_heading(f"Section {section + 1}: Research Overview", level=1)
        for p in range(8):
            doc.add_paragraph(
                f"S{section+1}/P{p+1}: "
                + "This research explores novel methodologies for automated document "
                  "processing using large language models and multi-agent frameworks. "
                  "Results demonstrate superior performance across all benchmark tasks. " * 2
            )
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_pdf() -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=LETTER)
    for page in range(3):
        c.setFont("Helvetica-Bold", 14)
        c.drawString(72, 720, f"Section {page + 1}: Results")
        c.setFont("Helvetica", 11)
        y = 690
        for line in range(20):
            c.drawString(72, y, f"Line {line}: The experimental results confirm our hypothesis conclusively.")
            y -= 18
        c.showPage()
    c.save()
    return buf.getvalue()


def make_text() -> bytes:
    content = """\
# My Research Paper

## Abstract

This paper presents a comprehensive study of NLP pipelines (Jones 2022).

## Introduction

Natural language processing has advanced rapidly in recent years [1][2].

### Background

Early work by Smith et al. (2019) established the baseline.

## Methods

We used Python 3.12 for all experiments.
The pipeline consists of five stages.

## Results

Accuracy improved by 12.3% over the baseline.

## References

Jones, A. (2022). NLP in the wild. *Journal of AI*. 10(2), 100-120.
Smith, B., & Lee, C. (2019). Baseline methods. *Proc. ACL*.
"""
    return content.encode("utf-8")


# ── Individual criterion runners ───────────────────────────────────────────────

async def sc1_all_formats_validate() -> None:
    """SC-1: All 3 formats produce IR objects that pass model_validate()."""
    print(f"\n{BOLD}SC-1  All formats → IR.model_validate(){RESET}")
    from app.agents.ingestion.docx_reader import DocxReader
    from app.agents.ingestion.pdf_reader import PdfReader
    from app.agents.ingestion.text_reader import TextReader
    from app.schemas.ir import IRSchema

    # ── .docx ──────────────────────────────────────────────────────────────────
    try:
        docx_bytes = make_docx()
        with patch("app.agents.ingestion.docx_reader.storage_service.upload_raw",
                   new_callable=AsyncMock, return_value="s3://key"):
            ir = await DocxReader().read(docx_bytes, "sc1_docx")
        validated = IRSchema.model_validate(ir.model_dump())
        passed = validated.source_format == "docx" and len(validated.elements) > 0
        record("SC-1a  .docx → IR validates", passed,
               f"source_format={validated.source_format}, elements={len(validated.elements)}")
    except Exception as exc:
        record("SC-1a  .docx → IR validates", False, str(exc))

    # ── .pdf ───────────────────────────────────────────────────────────────────
    try:
        pdf_bytes = make_pdf()
        with (
            patch("app.agents.ingestion.pdf_reader.storage_service.upload_raw",
                  new_callable=AsyncMock, return_value="s3://key"),
            patch.object(PdfReader, "_is_scanned", return_value=False),
        ):
            ir = await PdfReader().read(pdf_bytes, "sc1_pdf")
        validated = IRSchema.model_validate(ir.model_dump())
        passed = validated.source_format == "pdf" and len(validated.elements) > 0
        record("SC-1b  .pdf  → IR validates", passed,
               f"source_format={validated.source_format}, elements={len(validated.elements)}")
    except Exception as exc:
        record("SC-1b  .pdf  → IR validates", False, str(exc))

    # ── .txt / .md ─────────────────────────────────────────────────────────────
    try:
        txt_bytes = make_text()
        ir = await TextReader().read(txt_bytes, "sc1_txt")
        validated = IRSchema.model_validate(ir.model_dump())
        passed = validated.source_format == "txt" and len(validated.elements) > 0
        record("SC-1c  .txt  → IR validates", passed,
               f"source_format={validated.source_format}, elements={len(validated.elements)}")
    except Exception as exc:
        record("SC-1c  .txt  → IR validates", False, str(exc))


async def sc2_docx_50page_under_3s() -> None:
    """SC-2: 50-page .docx completes in < 3 seconds."""
    print(f"\n{BOLD}SC-2  50-page .docx < 3 seconds{RESET}")
    from app.agents.ingestion.docx_reader import DocxReader

    try:
        docx_bytes = make_50page_docx()
        page_count = len(docx_bytes)
        print(f"         Fixture size: {page_count / 1024:.1f} KB")

        with patch("app.agents.ingestion.docx_reader.storage_service.upload_raw",
                   new_callable=AsyncMock, return_value="s3://key"):
            t0 = time.perf_counter()
            ir = await DocxReader().read(docx_bytes, "sc2_50page")
            elapsed = time.perf_counter() - t0

        passed = elapsed < 3.0
        record(
            "SC-2   50-page .docx < 3 s",
            passed,
            f"Elapsed: {elapsed:.3f}s | Elements extracted: {len(ir.elements)} | Words: {ir.word_count}",
        )
    except Exception as exc:
        record("SC-2   50-page .docx < 3 s", False, str(exc))


async def sc3_scanned_pdf_triggers_ocr() -> None:
    """SC-3: Scanned PDF (image-only) triggers OcrReader delegation."""
    print(f"\n{BOLD}SC-3  Scanned PDF → OcrReader triggered{RESET}")
    from app.agents.ingestion.pdf_reader import PdfReader
    from app.agents.ingestion.ocr_reader import OcrReader
    from app.schemas.ir import IRSchema

    ocr_called = False

    async def fake_ocr_read(self, file_bytes: bytes, job_id: str) -> IRSchema:
        nonlocal ocr_called
        ocr_called = True
        return IRSchema(
            source_format="pdf",
            elements=[],
            word_count=0,
            metadata={"ocr_processed": True, "page_count": 1},
        )

    try:
        # Build fake fitz doc (low text density so _is_scanned returns True)
        fake_doc = MagicMock()
        fake_doc.page_count = 1
        fake_page = MagicMock()
        fake_page.rect.width = 595
        # Simulate image-only page: zero characters of text
        fake_page.get_text.return_value = ""
        fake_doc.__iter__ = MagicMock(return_value=iter([fake_page]))
        fake_doc.close = MagicMock()

        with (
            patch("app.agents.ingestion.pdf_reader.fitz.open", return_value=fake_doc),
            patch.object(OcrReader, "read", fake_ocr_read),
        ):
            reader = PdfReader()
            # _is_scanned uses real logic on our fake doc (text density = 0 < 0.5)
            ir = await reader.read(b"%PDF", "sc3_scanned")

        passed = ocr_called and ir.metadata.get("ocr_processed") is True
        record("SC-3   Scanned PDF → OcrReader", passed,
               f"ocr_called={ocr_called}, ocr_processed={ir.metadata.get('ocr_processed')}")
    except Exception as exc:
        record("SC-3   Scanned PDF → OcrReader", False, str(exc))


async def sc4_character_coverage() -> None:
    """SC-4: Extracted character count >= 95% of source document."""
    print(f"\n{BOLD}SC-4  Character coverage >= 95%{RESET}")
    from app.agents.ingestion.docx_reader import DocxReader
    from app.agents.ingestion.pdf_reader import PdfReader
    from app.agents.ingestion.text_reader import TextReader

    # ── .txt ────────────────────────────────────────────────────────────────────
    try:
        content = (
            "# Research Paper\n\n" +
            "The quick brown fox jumps over the lazy dog. " * 150 +
            "\n\n## References\n\nDoe (2022). Paper. Journal.\n"
        ).encode("utf-8")
        source_alpha = sum(1 for c in content.decode() if c.isalpha())
        ir = await TextReader().read(content, "sc4_txt")
        extracted = " ".join(e.raw_text for e in ir.elements)
        extracted_alpha = sum(1 for c in extracted if c.isalpha())
        coverage = extracted_alpha / source_alpha if source_alpha else 0
        passed = coverage >= 0.95
        record("SC-4a  .txt  coverage >= 95%", passed,
               f"Coverage: {coverage:.1%} ({extracted_alpha}/{source_alpha} alpha chars)")
    except Exception as exc:
        record("SC-4a  .txt  coverage >= 95%", False, str(exc))

    # ── .docx ───────────────────────────────────────────────────────────────────
    try:
        from docx import Document
        doc = Document()
        sample = "Natural language processing advances rapidly in academic research domains. " * 60
        for i in range(20):
            if i % 5 == 0:
                doc.add_heading(f"Chapter {i // 5 + 1}", level=1)
            else:
                doc.add_paragraph(f"Para {i}: " + sample)
        buf = BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        # Source alpha = all alpha chars in all paragraphs text
        doc2 = Document(BytesIO(docx_bytes))
        source_text = " ".join(p.text for p in doc2.paragraphs)
        source_alpha = sum(1 for c in source_text if c.isalpha())

        with patch("app.agents.ingestion.docx_reader.storage_service.upload_raw",
                   new_callable=AsyncMock, return_value="s3://key"):
            ir = await DocxReader().read(docx_bytes, "sc4_docx")

        extracted = " ".join(e.raw_text for e in ir.elements)
        extracted_alpha = sum(1 for c in extracted if c.isalpha())
        coverage = extracted_alpha / source_alpha if source_alpha else 0
        passed = coverage >= 0.95
        record("SC-4b  .docx coverage >= 95%", passed,
               f"Coverage: {coverage:.1%} ({extracted_alpha}/{source_alpha} alpha chars)")
    except Exception as exc:
        record("SC-4b  .docx coverage >= 95%", False, str(exc))

    # ── .pdf ────────────────────────────────────────────────────────────────────
    try:
        import fitz
        pdf_bytes = make_pdf()
        # Measure source: what fitz itself can read (native ground truth)
        raw_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        source_text = "".join(page.get_text() for page in raw_doc)
        raw_doc.close()
        source_alpha = sum(1 for c in source_text if c.isalpha())

        with (
            patch("app.agents.ingestion.pdf_reader.storage_service.upload_raw",
                  new_callable=AsyncMock, return_value="s3://key"),
            patch.object(PdfReader, "_is_scanned", return_value=False),
        ):
            ir = await PdfReader().read(pdf_bytes, "sc4_pdf")

        extracted = " ".join(e.raw_text for e in ir.elements)
        extracted_alpha = sum(1 for c in extracted if c.isalpha())
        coverage = extracted_alpha / source_alpha if source_alpha else 0
        passed = coverage >= 0.95
        record("SC-4c  .pdf  coverage >= 95%", passed,
               f"Coverage: {coverage:.1%} ({extracted_alpha}/{source_alpha} alpha chars)")
    except Exception as exc:
        record("SC-4c  .pdf  coverage >= 95%", False, str(exc))


def sc5_pytest_zero_failures() -> None:
    """SC-5: pytest app/agents/ingestion/tests/ passes with 0 failures."""
    print(f"\n{BOLD}SC-5  pytest app/agents/ingestion/tests/ — 0 failures{RESET}")
    result = subprocess.run(
        [
            sys.executable, "-m", "pytest",
            "app/agents/ingestion/tests/",
            "--asyncio-mode=auto",
            "-q",
            "--tb=short",
        ],
        capture_output=True,
        text=True,
    )
    output = result.stdout + result.stderr
    # Parse summary line
    lines = [l.strip() for l in output.splitlines() if l.strip()]
    summary = next((l for l in reversed(lines) if "passed" in l or "failed" in l or "error" in l), output[-200:])
    passed = result.returncode == 0 and "failed" not in summary and "error" not in summary.lower()
    record("SC-5   pytest 0 failures", passed, summary)
    if not passed:
        # Print last 30 lines for debugging
        for line in lines[-30:]:
            print(f"         {line}")


# ── Main ────────────────────────────────────────────────────────────────────────

async def main() -> None:
    print(f"\n{BOLD}{'='*70}{RESET}")
    print(f"{BOLD}  Agent Paperpal — DocIngestionAgent Success Criteria Validation{RESET}")
    print(f"{BOLD}{'='*70}{RESET}")

    await sc1_all_formats_validate()
    await sc2_docx_50page_under_3s()
    await sc3_scanned_pdf_triggers_ocr()
    await sc4_character_coverage()
    sc5_pytest_zero_failures()

    # ── Final report ───────────────────────────────────────────────────────────
    total  = len(results)
    passed = sum(1 for _, ok, _ in results if ok)
    failed = total - passed

    print(f"\n{BOLD}{'='*70}{RESET}")
    print(f"{BOLD}  FINAL REPORT  —  {passed}/{total} criteria PASSED{RESET}")
    print(f"{BOLD}{'='*70}{RESET}")
    for name, ok, detail in results:
        icon = f"{GREEN}✅{RESET}" if ok else f"{RED}❌{RESET}"
        print(f"  {icon}  {name}")

    if failed == 0:
        print(f"\n{GREEN}{BOLD}  🎉 ALL SUCCESS CRITERIA MET!{RESET}\n")
    else:
        print(f"\n{RED}{BOLD}  ⚠️  {failed} criterion/criteria FAILED — see details above.{RESET}\n")
        sys.exit(1)


if __name__ == "__main__":
    asyncio.run(main())
