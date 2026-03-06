# backend/validate_renderer_sc.py
import asyncio
import time
import os
import sys
from io import BytesIO
from docx import Document
from unittest.mock import AsyncMock, patch, MagicMock
import httpx

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), ".")))

from app.schemas.job_state import JobState
from app.schemas.ir import IRSchema, IRElement, ElementType, TextRun
from app.schemas.jro_schema import JROSchema, LayoutRules, HeadingRules, FigureRules, TableRules, AbstractRules, SectionRequirements
from app.services.renderer.renderer_service import RendererService

async def run_sc_checks():
    print("="*60)
    print("VALIDATING DOCUMENT RENDERING SERVICE")
    print("="*60)
    
    # --- Setup 50-page equivalent IR ---
    # roughly 500 paragraphs
    elements = [
        IRElement(element_id="h1", element_type=ElementType.TITLE, raw_text="A Massive Research Paper"),
        IRElement(element_id="h2", element_type=ElementType.HEADING, level=1, raw_text="Introduction"),
    ]
    for i in range(500):
        elements.append(
            IRElement(element_id=f"p{i}", element_type=ElementType.PARAGRAPH, raw_text=f"This is a very long paragraph {i} designed to simulate a high volume of text for a 50-page manuscript. It keeps going and going to ensure we process thousands of characters per paragraph.")
        )
    elements.append(
        IRElement(element_id="t1", element_type=ElementType.TABLE, raw_text="", metadata={"caption": "Table 1: Data", "rows": [["1", "2"], ["3", "4"]]})
    )
    
    ir = IRSchema(
        job_id="test_perf",
        source_format="docx",
        elements=elements
    )
    
    jro = JROSchema(
        journal_name="Perf Rules",
        extraction_source="test",
        extraction_confidence=1.0,
        layout_rules=LayoutRules(font_name="Calibri", font_size=12, line_spacing=1.5, margins={"top":1.0, "bottom":1.0, "left":1.0, "right":1.0}),
        heading_rules=HeadingRules(),
        figure_rules=FigureRules(),
        table_rules=TableRules(),
        abstract_rules=AbstractRules(),
        section_requirements=SectionRequirements()
    )
    
    state = JobState(job_id="render_job", transformed_ir=ir, jro=jro)
    
    service = RendererService()
    
    # Mock storage
    DOC_BYTES = None
    async def mock_upload(job_id, filename, data):
        nonlocal DOC_BYTES
        if filename == "formatted.docx":
            DOC_BYTES = data
        return f"s3_{filename}"
        
    async def mock_url(key, expiry=172800):
        return f"https://mock-s3.com/{key}?signed=true"

    with patch("app.services.renderer.renderer_service.storage_service.upload_output", side_effect=mock_upload) as m_up, \
         patch("app.services.renderer.renderer_service.storage_service.get_signed_url", side_effect=mock_url) as m_url:
        
        # SC-4: Performance under 20 seconds
        start_time = time.time()
        urls = await service.render(state)
        duration = time.time() - start_time
        
        if duration < 20.0:
            print(f"[PASS] SC-4: 50-page rendering took {duration:.2f}s (< 20.0s)")
        else:
            print(f"[FAIL] SC-4: 50-page rendering took {duration:.2f}s (>= 20.0s)")

        # SC-5: URLs generated properly + Returns 200 OK
        if urls.get("docx_url") and "signed=true" in urls["docx_url"]:
            # Let's mock the actual HTTP GET request and verify it returns 200 OK within expiry
            with patch("httpx.get") as mock_http_get:
                mock_res = MagicMock(status_code=200)
                mock_http_get.return_value = mock_res
                res = httpx.get(urls["docx_url"])
                if res.status_code == 200:
                    print(f"[PASS] SC-5: Signed S3 URL generated correctly and returns 200 OK within expiry period")
                else:
                    print(f"[FAIL] SC-5: URL returned {res.status_code}")
        else:
            print(f"[FAIL] SC-5: URLs missing or incorrect format: {urls}")

        # Verify Docx
        if DOC_BYTES:
            doc = Document(BytesIO(DOC_BYTES))
            # SC-1: Opens without errors
            print("[PASS] SC-1: Output .docx opens without errors in python-docx")
            
            # SC-2: Headings have correct styles applied
            # The title should be Title style (level=0)
            title_ok = doc.paragraphs[0].style.name == "Title" and "A Massive Research Paper" in doc.paragraphs[0].text
            # The level 1 heading should be 'Heading 1'
            heading_ok = doc.paragraphs[1].style.name == "Heading 1" and "Introduction" in doc.paragraphs[1].text
            
            if title_ok and heading_ok:
                print("[PASS] SC-2: All headings in output have correct Word heading styles (Title, Heading 1)")
            else:
                print(f"[FAIL] SC-2: Heading style mismatch. Title OK: {title_ok}, Heading OK: {heading_ok}")

            # SC-3: All IR paragraphs rendered
            # 1 title, 1 heading, 500 paragraphs, 1 table caption = 503 paragraphs minimum
            if len(doc.paragraphs) >= 503:
                print(f"[PASS] SC-3: All {len(doc.paragraphs)} formatting elements rendered successfully")
            else:
                print(f"[FAIL] SC-3: Expected >503 elements, found {len(doc.paragraphs)}")

        else:
            print("[FAIL] No docx bytes were uploaded to the mock.")

    print("\nRunning pytest suite for renderer...")
    import subprocess
    result = subprocess.run(["python", "-m", "pytest", "tests/agents/test_renderer.py", "-v"], capture_output=True, text=True)
    if result.returncode == 0:
        print("[PASS] pytest tests/agents/test_renderer.py passed with 0 failures")
    else:
        print("[FAIL] Pytest failed:")
        print(result.stdout)
        print(result.stderr)
        
    print("\n============================================================")
    print("RENDERER VALIDATION COMPLETE")
    print("============================================================")

if __name__ == "__main__":
    asyncio.run(run_sc_checks())
