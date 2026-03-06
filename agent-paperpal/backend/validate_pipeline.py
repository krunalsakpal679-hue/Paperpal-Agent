# backend/validate_pipeline.py
import asyncio
import time
from unittest.mock import AsyncMock, patch
from io import BytesIO

from app.schemas.job_state import JobState, JobStatus
from app.agents.ingestion.agent import DocIngestionAgent
from app.agents.parsing.agent import DocParseAgent
from app.agents.interpretation.agent import RuleInterpretAgent
from app.agents.transformation.agent import TransformAgent
from app.agents.validation.agent import ValidationAgent

def create_mock_docx() -> bytes:
    from docx import Document
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the intro paragraph. It has a citation (Smith, 2020).")
    doc.add_heading("Methods", level=1)
    doc.add_paragraph("We used standard methods to analyze the data.")
    doc.add_heading("References", level=1)
    doc.add_paragraph("Smith, J. (2020). A study on mock documents. Journal of Testing, 1(1), 1-10.")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

async def run_pipeline():
    print("="*60)
    print("PIPELINE END-TO-END VALIDATION")
    print("="*60)

    start_t = time.time()
    
    # Initialize state
    job_id = "e2e_test_job"
    state = JobState(
        job_id=job_id,
        metadata={
            "raw_s3_key": f"jobs/{job_id}/test.docx",
            "source_format": "docx",
            "journal_identifier": "apa",
            "style_guide_url": "https://www.nature.com/nature/for-authors/formatting-guide"
        }
    )
    
    docx_bytes = create_mock_docx()
    
    with patch("app.agents.ingestion.agent.storage_service.download_raw", new_callable=AsyncMock) as mock_dl_raw, \
         patch("app.agents.ingestion.docx_reader.storage_service.upload_raw", new_callable=AsyncMock) as mock_ul_raw, \
         patch("app.services.cache_service.cache_service.publish_progress", new_callable=AsyncMock) as mock_pub, \
         patch("app.agents.interpretation.scraper.JournalScraper.scrape", new_callable=AsyncMock) as mock_scrape, \
         patch("app.agents.ingestion.file_validator.magic.from_buffer") as mock_magic, \
         patch("app.services.cache_service.cache_service.get_jro", new_callable=AsyncMock) as mock_get_jro, \
         patch("app.services.cache_service.cache_service.set_jro", new_callable=AsyncMock) as mock_set_jro:
        
        mock_dl_raw.return_value = docx_bytes
        mock_ul_raw.return_value = "s3_key_image"
        mock_magic.return_value = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        mock_get_jro.return_value = None  # Simulate cache miss
        
        # Mock scrape to provide minimal text that the LLM interprets
        mock_scrape.return_value = ("Submission guidelines: We require APA citation style, abstract max 250 words, and figure captions below figures.", False)

        # Stage 1: Ingestion
        print("\n--- STAGE 1: INGESTION ---")
        agent1 = DocIngestionAgent()
        state = await agent1.run(state)
        print(f"Status: {state.status.value}")
        if state.errors: print("Errors:", [e.message for e in state.errors])
        if state.raw_ir: print(f"Raw IR extracted: {len(state.raw_ir.elements)} elements")

        # Stage 2: Parsing
        print("\n--- STAGE 2: PARSING ---")
        agent2 = DocParseAgent()
        state = await agent2.run(state)
        print(f"Status: {state.status.value}")
        if state.errors: print("Errors:", [e.message for e in state.errors])
        if state.annotated_ir: 
            print(f"Parsed Headings: {len(state.annotated_ir.metadata.get('headings_parsed', []))}")
            print(f"Parsed Citations: {len(state.annotated_ir.metadata.get('citations_parsed', []))}")
            print(f"Parsed References: {len(state.annotated_ir.metadata.get('references_parsed', []))}")

        # Stage 3: Interpretation
        print("\n--- STAGE 3: INTERPRETATION ---")
        agent3 = RuleInterpretAgent()
        state = await agent3.run(state)
        print(f"Status: {state.status.value}")
        if state.errors: print("Errors:", [e.message for e in state.errors])
        if state.jro: print(f"JRO extraction: {state.jro.journal_name} ({state.jro.citation_style})")

        # Stage 4: Transformation
        print("\n--- STAGE 4: TRANSFORMATION ---")
        agent4 = TransformAgent()
        state = await agent4.run(state)
        print(f"Status: {state.status.value}")
        if state.errors: print("Errors:", [e.message for e in state.errors])
        print(f"Changes applied: {len(state.change_log)}")

        # Stage 5: Validation
        print("\n--- STAGE 5: VALIDATION ---")
        agent5 = ValidationAgent()
        state = await agent5.run(state)
        print(f"Status: {state.status.value}")
        if state.errors: print("Errors:", [e.message for e in state.errors])
        if state.compliance_report:
            print(f"Validation Score: {state.compliance_report.overall_score}%")
            print(f"Total Issues: {state.compliance_report.total_issues}")
            for issue in state.compliance_report.items:
                print(f" - [{issue.severity.upper()}] {issue.description}")

    total_time = time.time() - start_t
    print("\n" + "="*60)
    print(f"END-TO-END VALIDATION COMPLETE (Time: {total_time:.2f}s)")
    if state.status == JobStatus.COMPLETED:
        print("PIPELINE SUCCESS!")
    else:
        print("PIPELINE FAILED!")
    print("="*60)

if __name__ == "__main__":
    asyncio.run(run_pipeline())
